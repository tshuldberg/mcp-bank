"""
Document format conversions: PDF, XLSX, DOCX.
All binary inputs are expected as base64-encoded strings.
"""

import base64
import io
import orjson
import polars as pl


def _decode_b64(data: str) -> bytes:
    return base64.b64decode(data)


# ─── PDF ──────────────────────────────────────────────────────────────────────

def pdf_to_text(data: str) -> str:
    import pymupdf  # fitz
    raw = _decode_b64(data)
    doc = pymupdf.open(stream=raw, filetype="pdf")
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n\n".join(pages)


def pdf_to_json(data: str) -> str:
    import pymupdf
    raw = _decode_b64(data)
    doc = pymupdf.open(stream=raw, filetype="pdf")
    pages = []
    for i, page in enumerate(doc):
        pages.append({
            "page": i + 1,
            "text": page.get_text(),
            "width": page.rect.width,
            "height": page.rect.height,
        })
    doc.close()
    return orjson.dumps({"pages": pages, "total_pages": len(pages)}).decode()


# ─── XLSX ─────────────────────────────────────────────────────────────────────

def xlsx_to_csv(data: str) -> str:
    raw = _decode_b64(data)
    df = pl.read_excel(io.BytesIO(raw))
    buf = io.StringIO()
    df.write_csv(buf)
    return buf.getvalue()


def xlsx_to_json(data: str) -> str:
    raw = _decode_b64(data)
    df = pl.read_excel(io.BytesIO(raw))
    return orjson.dumps(df.to_dicts()).decode()


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def docx_to_text(data: str) -> str:
    from docx import Document
    raw = _decode_b64(data)
    doc = Document(io.BytesIO(raw))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def docx_to_markdown(data: str) -> str:
    from docx import Document
    raw = _decode_b64(data)
    doc = Document(io.BytesIO(raw))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        style = para.style.name.lower()
        if style.startswith("heading 1"):
            lines.append(f"# {text}")
        elif style.startswith("heading 2"):
            lines.append(f"## {text}")
        elif style.startswith("heading 3"):
            lines.append(f"### {text}")
        elif style.startswith("heading"):
            lines.append(f"#### {text}")
        elif style.startswith("list"):
            lines.append(f"- {text}")
        else:
            lines.append(text)
    return "\n".join(lines)
